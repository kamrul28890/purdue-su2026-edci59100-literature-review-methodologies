"""Convert the Final Literature Review Manuscript .tex into a properly
formatted .docx, preserving numbering, citations, tables, and figures."""
import re
import pathlib

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = pathlib.Path(__file__).parent
SRC = ROOT.parent / "Final_Literature_Review_Manuscript_Summer26_Kamrul_MdKamruzzaman.tex"
FIGPNG = ROOT / "fig_png"
OUT = ROOT.parent / "Final_Literature_Review_Manuscript_Summer26_Kamrul_MdKamruzzaman.docx"

HEADING_BLUE = RGBColor(0x1F, 0x4E, 0x79)
LIGHT_BLUE = "D9EAF7"
SOFT_GRAY = "F4F6F8"
DARK_GRAY = RGBColor(0x40, 0x40, 0x40)

text = SRC.read_text(encoding="utf8")
body = text[text.index(r"\begin{document}") + len(r"\begin{document}"):text.index(r"\end{document}")]


# ---------------------------------------------------------------- utilities
def find_matching_brace(s, open_idx):
    """s[open_idx] must be '{'. Return index of matching '}'."""
    depth = 0
    i = open_idx
    while i < len(s):
        if s[i] == "{":
            depth += 1
        elif s[i] == "}":
            depth -= 1
            if depth == 0:
                return i
        i += 1
    raise ValueError("unbalanced braces")


def get_braced_arg(s, after_idx):
    """Given s and an index right after a command name, skip to the next
    '{' and return (content, end_index_after_closing_brace)."""
    start = s.index("{", after_idx)
    end = find_matching_brace(s, start)
    return s[start + 1:end], end + 1


# ----------------------------------------------------- citation numbering
bib_start = body.index(r"\begin{thebibliography}")
bib_block = body[bib_start:body.index(r"\end{thebibliography}")]
bibitem_re = re.compile(r"\\bibitem\{([^}]+)\}")
cite_keys_in_order = bibitem_re.findall(bib_block)
cite_num = {k: i + 1 for i, k in enumerate(cite_keys_in_order)}

# ----------------------------------------------------- figure/table/section numbering
fig_labels = re.findall(r"\\begin\{figure\}.*?\\label\{(fig:[a-z]+)\}", body, re.S)
fig_num = {lbl: i + 1 for i, lbl in enumerate(fig_labels)}

# tables: appear either as \begin{table}[H] ... \label{tab:xxx} or inside \begin{landscape}\begin{longtable}
table_label_re = re.compile(r"\\label\{(tab:[a-z]+)\}")
table_labels = []
for m in re.finditer(r"\\caption\{.*?\\label\{(tab:[a-z]+)\}", body, re.S):
    table_labels.append(m.group(1))
table_num = {lbl: i + 1 for i, lbl in enumerate(table_labels)}

# sections/subsections: build label map by walking sequential heading + label pairs (sec: labels)
sec_label_num = {}


# ---------------------------------------------------------- inline run model
class Run:
    __slots__ = ("text", "bold", "italic", "superscript", "url", "anchor")

    def __init__(self, text, bold=False, italic=False, superscript=False, url=None, anchor=None):
        self.text = text
        self.bold = bold
        self.italic = italic
        self.superscript = superscript
        self.url = url
        self.anchor = anchor


def resolve_ref(label):
    if label in fig_num:
        return str(fig_num[label])
    if label in table_num:
        return str(table_num[label])
    if label in sec_label_num:
        return sec_label_num[label]
    return "?"


def clean_basic_text(s):
    s = s.replace(r"\%", "%").replace(r"\&", "&").replace(r"\_", "_")
    s = s.replace(r"\textbar{}", "|").replace(r"\textbar", "|")
    s = s.replace(r"\kappa", "κ")
    s = s.replace("{,}", ",")
    s = s.replace("$", "")
    s = re.sub(r"\\(small|footnotesize|scriptsize|tiny|normalsize|Large|large|bfseries|itshape|centering|noindent)\b", "", s)
    s = s.replace("---", "—").replace("--", "–")
    s = s.replace("``", "“").replace("''", "”")
    s = s.replace("~", " ")
    s = s.replace(r"\ ", " ")
    s = s.replace(r"\\", " ")
    s = re.sub(r"\\textwidth", "", s)
    s = s.replace(r"\bullet", "•")
    return s


def parse_inline(s, bold=False, italic=False):
    """Recursively parse a LaTeX inline string into a flat list of Run objects."""
    runs = []
    i = 0
    n = len(s)
    buf = ""

    def flush():
        nonlocal buf
        if buf:
            runs.append(Run(clean_basic_text(buf), bold=bold, italic=italic))
            buf = ""

    while i < n:
        ch = s[i]
        if ch == "\\":
            m = re.match(r"\\(textbf|textit|emph|cite|ref|href|textsuperscript)\b", s[i:])
            if m:
                cmd = m.group(1)
                after = i + m.end()
                if cmd == "textbf":
                    flush()
                    content, end = get_braced_arg(s, after)
                    runs.extend(parse_inline(content, bold=True, italic=italic))
                    i = end
                    continue
                if cmd in ("textit", "emph"):
                    flush()
                    content, end = get_braced_arg(s, after)
                    runs.extend(parse_inline(content, bold=bold, italic=True))
                    i = end
                    continue
                if cmd == "textsuperscript":
                    flush()
                    content, end = get_braced_arg(s, after)
                    runs.append(Run(clean_basic_text(content), bold=bold, italic=italic, superscript=True))
                    i = end
                    continue
                if cmd == "cite":
                    flush()
                    content, end = get_braced_arg(s, after)
                    keys = [k.strip() for k in content.split(",")]
                    runs.append(Run("[", bold=bold, italic=italic))
                    for kidx, k in enumerate(keys):
                        if kidx > 0:
                            runs.append(Run(",", bold=bold, italic=italic))
                        num = cite_num.get(k, "?")
                        runs.append(Run(str(num), bold=bold, italic=italic, anchor=f"ref{num}"))
                    runs.append(Run("]", bold=bold, italic=italic))
                    i = end
                    continue
                if cmd == "ref":
                    flush()
                    content, end = get_braced_arg(s, after)
                    runs.append(Run(resolve_ref(content.strip()), bold=bold, italic=italic))
                    i = end
                    continue
                if cmd == "href":
                    flush()
                    url, end1 = get_braced_arg(s, after)
                    label, end2 = get_braced_arg(s, end1)
                    runs.append(Run(clean_basic_text(label), bold=bold, italic=italic, url=url))
                    i = end2
                    continue
            # textbar{} / kappa / bullet / other single tokens handled by clean_basic_text via buf
            m2 = re.match(r"\\[a-zA-Z]+\*?\{?", s[i:])
            if m2:
                buf += m2.group(0)
                i += m2.end()
                continue
        buf += ch
        i += 1
    flush()
    return runs


def add_runs_to_paragraph(p, runs, font_size=11, font_name="Times New Roman"):
    for r in runs:
        if not r.text:
            continue
        run = p.add_run(r.text)
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.bold = r.bold
        run.italic = r.italic
        if r.superscript:
            run.font.superscript = True
        if r.url:
            wrap_run_as_link(p, run, url=r.url, style_as_link=True)
        elif r.anchor:
            wrap_run_as_link(p, run, anchor=r.anchor, style_as_link=False)


def wrap_run_as_link(paragraph, run, url=None, anchor=None, style_as_link=True):
    """Convert an existing run into a real Word hyperlink: external (url, via
    relationship id) or internal (anchor, jumps to a same-document bookmark)."""
    hyperlink = OxmlElement("w:hyperlink")
    if url is not None:
        part = paragraph.part
        r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
        hyperlink.set(qn("r:id"), r_id)
    else:
        hyperlink.set(qn("w:anchor"), anchor)
    run._r.getparent().remove(run._r)
    hyperlink.append(run._r)
    if style_as_link:
        rPr = run._r.find(qn("w:rPr"))
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            run._r.insert(0, rPr)
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "1565C0")
        rPr.append(color)
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rPr.append(u)
    paragraph._p.append(hyperlink)


_bookmark_id_seq = [0]


def add_bookmark_around_run(run, name):
    _bookmark_id_seq[0] += 1
    bid = _bookmark_id_seq[0]
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bid))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bid))
    run._r.addprevious(start)
    run._r.addnext(end)


print(f"Citations mapped: {len(cite_num)}; figures: {fig_num}; tables found: {len(table_num)}")

# ---------------------------------------------------------- section numbering pre-pass
heading_re = re.compile(r"\\(section|subsection|subsubsection)\{")
sec_ctr, sub_ctr, subsub_ctr = 0, 0, 0
headings = []  # list of dicts: idx, level, number, title_raw, end_idx_after_title_brace
for m in heading_re.finditer(body):
    level = m.group(1)
    title, end = get_braced_arg(body, m.end() - 1)
    if level == "section":
        sec_ctr += 1
        sub_ctr = 0
        subsub_ctr = 0
        number = f"{sec_ctr}"
    elif level == "subsection":
        sub_ctr += 1
        subsub_ctr = 0
        number = f"{sec_ctr}.{sub_ctr}"
    else:
        subsub_ctr += 1
        number = f"{sec_ctr}.{sub_ctr}.{subsub_ctr}"
    # check for an immediately following \label{sec:...}, possibly after a comment line
    after = body[end:end + 90]
    lm = re.match(r"\s*(?:%[^\n]*\n\s*)?\\label\{(sec:[a-z]+)\}", after)
    label = None
    label_end = end
    if lm:
        label = lm.group(1)
        sec_label_num[label] = number
        label_end = end + lm.end()
    headings.append({
        "start": m.start(), "level": level, "number": number,
        "title": title, "title_end": label_end, "label": label,
    })

for h in headings:
    print(h["number"], h["level"], h["title"][:50], "label=", h["label"])

# ---------------------------------------------------------- special-block spans
intro_start = headings[0]["start"]  # \section{Introduction}
bib_section_start = body.index(r"\begin{thebibliography}")

spans = []  # dicts: start, end, kind, raw
for h in headings:
    spans.append({"start": h["start"], "end": h["title_end"], "kind": "heading", "h": h})

table_pat = re.compile(r"\\begin\{table\}\[H\].*?\\end\{table\}", re.S)
for m in table_pat.finditer(body, intro_start, bib_section_start):
    spans.append({"start": m.start(), "end": m.end(), "kind": "table", "raw": m.group(0)})

landscape_pat = re.compile(r"\\begin\{landscape\}.*?\\end\{landscape\}", re.S)
for m in landscape_pat.finditer(body, intro_start, bib_section_start):
    spans.append({"start": m.start(), "end": m.end(), "kind": "table", "raw": m.group(0)})

figure_pat = re.compile(r"\\begin\{figure\}\[H\].*?\\end\{figure\}", re.S)
for m in figure_pat.finditer(body, intro_start, bib_section_start):
    spans.append({"start": m.start(), "end": m.end(), "kind": "figure", "raw": m.group(0)})

itemize_pat = re.compile(r"\\begin\{itemize\}.*?\\end\{itemize\}", re.S)
for m in itemize_pat.finditer(body, intro_start, bib_section_start):
    spans.append({"start": m.start(), "end": m.end(), "kind": "itemize", "raw": m.group(0)})

quote_pat = re.compile(r"\\begin\{quote\}.*?\\end\{quote\}", re.S)
for m in quote_pat.finditer(body, intro_start, bib_section_start):
    spans.append({"start": m.start(), "end": m.end(), "kind": "quote", "raw": m.group(0)})

spans.append({"start": bib_section_start, "end": len(body), "kind": "bibliography", "raw": body[bib_section_start:]})

spans.sort(key=lambda d: d["start"])


# =================================================================
#  DOCUMENT BUILD
# =================================================================
doc = Document()

# --- base styles ---
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(12)
normal.paragraph_format.line_spacing = 2.0
normal.paragraph_format.space_after = Pt(0)
normal.paragraph_format.first_line_indent = Inches(0.5)

sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = Inches(1)
sec.left_margin = sec.right_margin = Inches(1)


def set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "BFBFBF")
        borders.append(el)
    tblPr.append(borders)


def start_landscape_section():
    new_sec = doc.add_section(WD_SECTION.NEW_PAGE)
    new_sec.orientation = WD_ORIENT.LANDSCAPE
    new_sec.page_width, new_sec.page_height = Inches(11), Inches(8.5)
    new_sec.left_margin = new_sec.right_margin = Inches(0.6)
    new_sec.top_margin = new_sec.bottom_margin = Inches(0.75)
    return new_sec


def end_landscape_section():
    new_sec = doc.add_section(WD_SECTION.NEW_PAGE)
    new_sec.orientation = WD_ORIENT.PORTRAIT
    new_sec.page_width, new_sec.page_height = Inches(8.5), Inches(11)
    new_sec.left_margin = new_sec.right_margin = Inches(1)
    new_sec.top_margin = new_sec.bottom_margin = Inches(1)
    return new_sec


def compute_column_widths_in(cols, target_width_in):
    target_cm = target_width_in * 2.54
    raw_cm = []
    for c in cols:
        m = re.match(r"L\{([\d.]+)cm\}", c)
        raw_cm.append(float(m.group(1)) if m else None)
    known_sum = sum(w for w in raw_cm if w is not None)
    n_unknown = sum(1 for w in raw_cm if w is None)
    leftover = max(target_cm - known_sum, n_unknown * 1.8 if n_unknown else 0)
    y_width = leftover / n_unknown if n_unknown else 0
    cm_widths = [w if w is not None else y_width for w in raw_cm]
    scale = target_cm / sum(cm_widths)
    return [Inches((w * scale) / 2.54) for w in cm_widths]


def set_column_widths(table, widths_in):
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for cidx, cell in enumerate(row.cells):
            cell.width = widths_in[cidx]
    for cidx, col in enumerate(table.columns):
        col.width = widths_in[cidx]


def repeat_header_row(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader")
    th.set(qn("w:val"), "true")
    trPr.append(th)


def add_heading_paragraph(number, level, title_raw):
    style_name = {"section": "Heading 1", "subsection": "Heading 2", "subsubsection": "Heading 3"}[level]
    p = doc.add_paragraph(style=style_name)
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(14 if level == "section" else 8)
    p.paragraph_format.space_after = Pt(6)
    runs = parse_inline(title_raw)
    full_text = number + ". " + "".join(r.text for r in runs)
    run = p.add_run(full_text)
    run.font.name = "Times New Roman"
    run.font.bold = True
    run.font.size = Pt(15 if level == "section" else (13 if level == "subsection" else 12))
    run.font.color.rgb = HEADING_BLUE if level != "subsubsection" else DARK_GRAY
    return p


def add_body_paragraph(raw_text, align=None, indent=True, font_size=12, space_after=6, line_spacing=2.0):
    runs = parse_inline(raw_text)
    if not any(r.text.strip() for r in runs):
        return None
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_after = Pt(space_after)
    if not indent:
        p.paragraph_format.first_line_indent = Inches(0)
    if align is not None:
        p.alignment = align
    add_runs_to_paragraph(p, runs, font_size=font_size)
    return p


COMMENT_LINE_RE = re.compile(r"^\s*%.*$", re.M)


def strip_comments(s):
    return COMMENT_LINE_RE.sub("", s)


def split_paragraphs(raw):
    raw = strip_comments(raw)
    chunks = re.split(r"\n\s*\n", raw)
    return [c.strip() for c in chunks if c.strip()]


def emit_free_text(raw):
    for chunk in split_paragraphs(raw):
        add_body_paragraph(chunk)


# ------------------------------------------------------- table cell column parsing
def parse_colspec(colspec):
    colspec = colspec.replace("|", " ")
    tokens = re.findall(r"L\{[^}]*\}|Y|[lcr]|p\{[^}]*\}", colspec)
    return tokens


def split_row_cells(row_text):
    """Split a table row on unescaped & characters."""
    cells = re.split(r"(?<!\\)&", row_text)
    return [c.strip() for c in cells]


def render_table(raw, landscape=False):
    cap_m = re.search(r"\\caption\{", raw)
    caption_text, after_cap = get_braced_arg(raw, cap_m.end() - 1)
    label_m = re.search(r"\\label\{(tab:[a-z]+)\}", raw)
    label = label_m.group(1)
    tnum = table_num[label]

    anchor_m = re.search(r"\\begin\{tabularx\}\{\\textwidth\}|\\begin\{longtable\}", raw)
    is_longtable = "longtable" in anchor_m.group(0)
    colspec_str, body_start = get_braced_arg(raw, anchor_m.end())
    cols = parse_colspec(colspec_str)
    ncols = len(cols)
    end_kw = r"\end{longtable}" if is_longtable else r"\end{tabularx}"
    table_body = raw[body_start:raw.index(end_kw)]
    table_body = strip_comments(table_body)

    # drop repeated-header longtable machinery: keep only the firsthead block content as the header,
    # drop everything between \endfirsthead/\endhead markers used for continuation pages.
    table_body = re.sub(r"\\endfirsthead.*?\\endhead", "", table_body, flags=re.S)
    table_body = re.sub(r"\\endfoot.*?\\endlastfoot", "", table_body, flags=re.S)

    # caption already declared inline for longtable; strip any stray \caption{...}\\ \label{...}\\ inside
    table_body = re.sub(r"\\caption\{.*?\\label\{tab:[a-z]+\}\\\\", "", table_body, flags=re.S)

    raw_rows = [r for r in table_body.split(r"\\") if r.strip()]
    parsed_rows = []  # list of (cells:list[str], is_rule_only:bool, is_multicolumn:tuple|None)
    for r in raw_rows:
        r = re.sub(r"\\(top|mid|bottom)rule\b", "", r)
        r = re.sub(r"\\rowcolors\{[^}]*\}\{[^}]*\}\{[^}]*\}", "", r)
        r = re.sub(r"\\rowcolor\{[^}]*\}", "", r)
        r = r.strip()
        if not r:
            continue
        mcm = re.match(r"\\multicolumn\{(\d+)\}\{[lcr]\}\{(.*)\}$", r, re.S)
        if mcm:
            span = int(mcm.group(1))
            parsed_rows.append({"multicolumn": span, "text": mcm.group(2)})
            continue
        cells = split_row_cells(r)
        parsed_rows.append({"cells": cells})

    if landscape:
        start_landscape_section()
        target_width_in = 9.7
    else:
        target_width_in = 6.5

    table = doc.add_table(rows=0, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    col_widths = compute_column_widths_in(cols, target_width_in)

    header_done = False
    for ridx, row in enumerate(parsed_rows):
        wrow = table.add_row()
        if "multicolumn" in row:
            cell0 = wrow.cells[0]
            for extra in range(1, ncols):
                cell0 = cell0.merge(wrow.cells[extra])
            set_cell_shading(cell0, "EAF2FB")
            p = cell0.paragraphs[0]
            p.paragraph_format.first_line_indent = Inches(0)
            p.paragraph_format.line_spacing = 1.0
            runs = parse_inline(row["text"], bold=True)
            add_runs_to_paragraph(p, runs, font_size=10)
            continue
        cells = row["cells"]
        for cidx in range(ncols):
            ctext = cells[cidx] if cidx < len(cells) else ""
            ctext = ctext.strip()
            if ctext in ("~", "-"):
                ctext = ""
            cell = wrow.cells[cidx]
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = Inches(0)
            p.paragraph_format.line_spacing = 1.0
            is_header = not header_done
            runs = parse_inline(ctext, bold=is_header)
            add_runs_to_paragraph(p, runs, font_size=10 if not is_header else 10)
            if is_header:
                set_cell_shading(cell, LIGHT_BLUE)
            elif ridx % 2 == 0:
                set_cell_shading(cell, SOFT_GRAY)
        if not header_done:
            repeat_header_row(wrow)
            header_done = True

    set_column_widths(table, col_widths)

    cap_p = doc.add_paragraph()
    cap_p.paragraph_format.first_line_indent = Inches(0)
    cap_p.paragraph_format.space_before = Pt(4)
    cap_p.paragraph_format.space_after = Pt(10)
    cap_p.paragraph_format.line_spacing = 1.0
    run = cap_p.add_run(f"Table {tnum}. ")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Times New Roman"
    add_runs_to_paragraph(cap_p, parse_inline(caption_text), font_size=10)

    if landscape:
        end_landscape_section()


def render_figure(raw):
    cap_m = re.search(r"\\caption\{", raw)
    caption_text, _ = get_braced_arg(raw, cap_m.end() - 1)
    label_m = re.search(r"\\label\{(fig:[a-z]+)\}", raw)
    label = label_m.group(1)
    fnum = fig_num[label]
    img_name = label.split(":")[1] + "-1.png"
    img_path = FIGPNG / img_name

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run()
    run.add_picture(str(img_path), width=Inches(6.2))

    cap_p = doc.add_paragraph()
    cap_p.paragraph_format.first_line_indent = Inches(0)
    cap_p.paragraph_format.space_after = Pt(10)
    cap_p.paragraph_format.line_spacing = 1.0
    run = cap_p.add_run(f"Figure {fnum}. ")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = "Times New Roman"
    add_runs_to_paragraph(cap_p, parse_inline(caption_text), font_size=10)


def render_itemize(raw):
    inner = re.search(r"\\begin\{itemize\}(.*)\\end\{itemize\}", raw, re.S).group(1)
    items = re.split(r"\\item\b", inner)[1:]
    for it in items:
        it = strip_comments(it).strip()
        if not it:
            continue
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing = 2.0
        p.paragraph_format.first_line_indent = Inches(0)
        add_runs_to_paragraph(p, parse_inline(it), font_size=12)


def render_quote(raw):
    inner = re.search(r"\\begin\{quote\}(.*)\\end\{quote\}", raw, re.S).group(1)
    for chunk in split_paragraphs(inner):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(8)
        add_runs_to_paragraph(p, parse_inline(chunk), font_size=11)


def render_bibliography(raw):
    add_heading_paragraph_plain("References")
    items = re.split(r"\\bibitem\{[^}]+\}", raw)[1:]
    for i, it in enumerate(items, start=1):
        it = strip_comments(it).strip()
        it = re.sub(r"\\end\{thebibliography\}.*$", "", it, flags=re.S).strip()
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(f"[{i}] ")
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        add_bookmark_around_run(run, f"ref{i}")
        add_runs_to_paragraph(p, parse_inline(it), font_size=11)


def add_heading_paragraph_plain(text):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = HEADING_BLUE


print("\n--- BUILD: front matter + body ---")


def extract_top_level_braced_groups(s):
    groups = []
    i = 0
    n = len(s)
    while i < n:
        if s[i] == "{":
            end = find_matching_brace(s, i)
            groups.append(s[i + 1:end])
            i = end + 1
        else:
            i += 1
    return groups


NOARG_FMT = ("Large", "large", "normalsize", "small", "footnotesize", "scriptsize", "tiny", "bfseries", "itshape", "em")


def strip_leading_format_commands(s):
    s = s.lstrip()
    changed = True
    while changed:
        changed = False
        for cmd in NOARG_FMT:
            pat = "\\" + cmd
            if s.startswith(pat):
                s = s[len(pat):].lstrip()
                changed = True
        m = re.match(r"\\color\{[^}]*\}", s)
        if m:
            s = s[m.end():].lstrip()
            changed = True
    return s


# ------------------------------------------------------------- front matter
front_raw = body[:intro_start]
center_m = re.search(r"\\begin\{center\}(.*?)\\end\{center\}", front_raw, re.S)
title_groups = extract_top_level_braced_groups(center_m.group(1))
title_groups = [strip_leading_format_commands(g) for g in title_groups]

title_text, author_line, affil_line, corr_line, course_line, discipline_line = title_groups

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.first_line_indent = Inches(0)
p.paragraph_format.line_spacing = 1.15
p.paragraph_format.space_after = Pt(14)
run = p.add_run(title_text)
run.font.name = "Times New Roman"
run.font.size = Pt(18)
run.bold = True
run.font.color.rgb = HEADING_BLUE

for line, size in ((author_line, 13), (affil_line, 11), (corr_line, 11), (course_line, 11), (discipline_line, 11)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(4)
    add_runs_to_paragraph(p, parse_inline(line), font_size=size)

doc.add_page_break()

# ------------------------------------------------------------- abstract + keywords
abs_center_m = list(re.finditer(r"\\begin\{center\}(.*?)\\end\{center\}", front_raw, re.S))[1]
mp_inner = re.search(r"\\begin\{minipage\}\{[^}]*\}(.*?)\\end\{minipage\}", abs_center_m.group(1), re.S).group(1)
abs_groups = extract_top_level_braced_groups(mp_inner)
abstract_label = strip_leading_format_commands(abs_groups[0])
abstract_body = strip_leading_format_commands(abs_groups[1])
keywords_group = strip_leading_format_commands(abs_groups[2])

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Inches(0)
p.paragraph_format.space_after = Pt(8)
p.paragraph_format.line_spacing = 1.0
run = p.add_run(abstract_label.strip())
run.bold = True
run.font.name = "Times New Roman"
run.font.size = Pt(13)
run.font.color.rgb = HEADING_BLUE

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Inches(0)
p.paragraph_format.line_spacing = 1.5
p.paragraph_format.space_after = Pt(10)
add_runs_to_paragraph(p, parse_inline(abstract_body.strip()), font_size=11)

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Inches(0)
p.paragraph_format.line_spacing = 1.5
p.paragraph_format.space_after = Pt(10)
add_runs_to_paragraph(p, parse_inline(keywords_group.strip()), font_size=11)

# ------------------------------------------------------------- main body walk
cursor = intro_start
for sp in spans:
    if sp["start"] > cursor:
        emit_free_text(body[cursor:sp["start"]])
    if sp["kind"] == "heading":
        h = sp["h"]
        add_heading_paragraph(h["number"], h["level"], h["title"])
        cursor = h["title_end"]
    elif sp["kind"] == "table":
        landscape = "landscape" in sp["raw"][:25]
        render_table(sp["raw"], landscape=landscape)
        cursor = sp["end"]
    elif sp["kind"] == "figure":
        render_figure(sp["raw"])
        cursor = sp["end"]
    elif sp["kind"] == "itemize":
        render_itemize(sp["raw"])
        cursor = sp["end"]
    elif sp["kind"] == "quote":
        render_quote(sp["raw"])
        cursor = sp["end"]
    elif sp["kind"] == "bibliography":
        render_bibliography(sp["raw"])
        cursor = sp["end"]

doc.save(str(OUT))
print(f"\nSaved: {OUT}")
print(f"Paragraphs: {len(doc.paragraphs)}  Tables: {len(doc.tables)}")




